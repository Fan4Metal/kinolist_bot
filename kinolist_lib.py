import argparse
import glob
import io
import logging
import os
import re
import sys
import textwrap
import time
from copy import deepcopy
from pathlib import Path

import requests
from docx import Document
from docx.shared import Cm, Pt, RGBColor
from kinopoisk_unofficial.kinopoisk_api_client import KinopoiskApiClient
from kinopoisk_unofficial.request.films.film_request import FilmRequest
from kinopoisk_unofficial.request.staff.staff_request import StaffRequest
from mutagen.mp4 import MP4, MP4Cover
from PIL import Image
from tqdm import tqdm

LIB_VER = "0.2.7"

log = logging.getLogger("Lib")


def is_api_ok(api):
    '''Проверка авторизации.'''
    try:
        api_client = KinopoiskApiClient(api)
        request = FilmRequest(328)
        api_client.films.send_film_request(request)
    except Exception:
        return False
    else:
        return True


def image_to_file(image):
    """Return `image` as PNG file-like object."""
    image_file = io.BytesIO()
    image.save(image_file, format="PNG")
    return image_file


def get_resource_path(relative_path):
    '''
    Определение пути для запуска из автономного exe файла.

    Pyinstaller cоздает временную папку, путь в _MEIPASS.
    '''
    try:
        base_path = sys._MEIPASS
    except Exception:
        base_path = os.path.abspath(".")
    return os.path.join(base_path, relative_path)


def copy_table_after(table, paragraph):
    '''Копирование таблицы в указанный параграф.'''
    tbl, p = table._tbl, paragraph._p
    new_tbl = deepcopy(tbl)
    p.addnext(new_tbl)


def clone_first_table(document: Document, num):
    '''Клонирует первую таблицу в документе num раз.'''
    template = document.tables[0]
    paragraph = document.paragraphs[0]
    for i in range(num):
        copy_table_after(template, paragraph)
        paragraph = document.add_paragraph()


def find_kp_id_in_title(title: str):
    id = re.search(r"KP~(\d+)", title)
    if id:
        return id.group(1)


def find_kp_id(film_list, api):
    """Gets list of kinopoisk ids for list of films

    Args:
        film_list (list): List of movie titles for search
        api (string): Kinopoisk API token

    Returns:
        list: List of two elements:
                 0. list of found kinopoisk ids
                 1. list of items that have not been found
    """
    film_codes = []
    film_not_found = []
    for film in film_list:

        code_in_name = find_kp_id_in_title(film)
        if code_in_name:
            try:
                film_info = get_film_info(code_in_name, api)
                log.info(f'Найден фильм: {film_info[0]} ({film_info[1]}), kinopoisk id: {code_in_name}')
                film_codes.append(code_in_name)
                continue
            except Exception:
                film_not_found.append(code_in_name)
                continue
        time.sleep(0.2)
        payload = {'keyword': film, 'page': 1}
        headers = {'X-API-KEY': api, 'Content-Type': 'application/json'}
        try:
            r = requests.get('https://kinopoiskapiunofficial.tech/api/v2.1/films/search-by-keyword',
                             headers=headers,
                             params=payload)
            if r.status_code == 200:
                resp_json = r.json()
                if resp_json['searchFilmsCountResult'] == 0:
                    log.info(f'{film} не найден')
                    film_not_found.append(film)
                    continue
                else:
                    id = resp_json['films'][0]['filmId']
                    year = resp_json['films'][0]['year']
                    if 'nameRu' in resp_json['films'][0]:
                        found_film = resp_json['films'][0]['nameRu']
                    else:
                        found_film = resp_json['films'][0]['nameEn']
                    log.info(f'Найден фильм: {found_film} ({year}), kinopoisk id: {id}')
                    film_codes.append(id)
            else:
                log.warning('Ошибка доступа к https://kinopoiskapiunofficial.tech')
                return
        except Exception as e:
            log.warning("Exeption:", str(e))
            log.info(f'{film} не найден (exeption)')
            film_not_found.append(film)
            continue
    result = []
    result.append(film_codes)
    result.append(film_not_found)
    return result


def get_film_info(film_code, api, shorten=False):
    '''
    Получение информации о фильме с помощью kinopoisk_api_client.

            Элементы списка:
                0 - название фильма на русском языке
                1 - год
                2 - рейтинг Кинопоиска
                3 - список стран
                4 - описание
                5 - ссылка на постер
                6 - имя файла без расширения
                7 - режиссер
             8:17 - 10 актеров
               18 - Постер размером 360x540 в формате PIL.Image.Image
    '''
    api_client = KinopoiskApiClient(api)
    request_staff = StaffRequest(film_code)
    response_staff = api_client.staff.send_staff_request(request_staff)
    staff_list = []
    if len(response_staff.items) >= 11:
        for i in range(0, 11):  # загружаем 11 персоналий (режиссер + 10 актеров)
            if response_staff.items[i].name_ru == '':
                staff_list.append(response_staff.items[i].name_en)
            else:
                staff_list.append(response_staff.items[i].name_ru)
    else:
        for i in range(0, len(response_staff.items)):
            if response_staff.items[i].name_ru == '':
                staff_list.append(response_staff.items[i].name_en)
            else:
                staff_list.append(response_staff.items[i].name_ru)
        for i in range(11 - len(response_staff.items)):
            staff_list.append("")
    request_film = FilmRequest(film_code)
    response_film = api_client.films.send_film_request(request_film)
    # с помощью регулярного выражения находим значение стран в кавычках ''
    countries = re.findall("'([^']*)'", str(response_film.film.countries))
    # имя файла
    if response_film.film.name_ru is not None:
        file_name = response_film.film.name_ru
        film_name = response_film.film.name_ru
    else:
        file_name = response_film.film.name_original
        film_name = response_film.film.name_original
    # очистка имени файла от запрещенных символов
    trtable = file_name.maketrans('', '', '\/:*?"<>')
    file_name = file_name.translate(trtable)

    # Сокращение описания фильма
    if shorten:
        description = response_film.film.description.replace("\n\n", " ")
        description = textwrap.shorten(description,
                                       665,
                                       fix_sentence_endings=True,
                                       break_long_words=False,
                                       placeholder='...')
    else:
        description = response_film.film.description

    film_list = [
        film_name, response_film.film.year, response_film.film.rating_kinopoisk, countries, description,
        response_film.film.poster_url, file_name
    ]
    result = film_list + staff_list
    # загрузка постера
    cover_url = response_film.film.poster_url
    cover = requests.get(cover_url, stream=True)
    if cover.status_code == 200:
        cover.raw.decode_content = True
        image = Image.open(cover.raw)
        width, height = image.size
        # обрезка до соотношения сторон 1x1.5
        if width > (height / 1.5):
            image = image.crop((((width - height / 1.5) / 2), 0, ((width - height / 1.5) / 2) + height / 1.5, height))
        image.thumbnail((360, 540))
        rgb_image = image.convert('RGB')  # Fix "OSError: cannot write mode RGBA as JPEG"
        result.append(rgb_image)
    else:
        result.append("")
    return result


def get_full_film_list(film_codes: list, api: str, shorten=False):
    """Загружает информацию о фильмах

    Args:
        film_codes (list): Список kinopoisk_id фильмов
        api (str): Kinopoisk API token
        shorten (boolean): Option to shorten movie descriptions
    Returns:
        list: Список с полной информацией о фильмах для записи в таблицу.
    """
    full_films_list = []
    for film_code in tqdm(film_codes, desc="Загрузка информации...   "):
        try:
            film_info = get_film_info(film_code, api, shorten)
            full_films_list.append(film_info)
        except Exception as e:
            log.warning("Exeption:", str(e))
        else:
            continue
    return full_films_list


def write_film_to_table(current_table, filminfo: list):
    """Заполнение таблицы в файле docx.

    Args:
        current_table (Document object loaded from *docx*): указатель на текущую таблицу
        filminfo (list): информация о фильме
    """
    paragraph = current_table.cell(0, 1).paragraphs[0]  # название фильма + рейтинг
    if filminfo[2] == None:
        run = paragraph.add_run(str(filminfo[0]) + ' - ' + 'нет рейтинга')
    else:
        run = paragraph.add_run(str(filminfo[0]) + ' - ' + 'Кинопоиск ' + str(filminfo[2]))
    run.font.name = 'Arial'
    run.font.size = Pt(11)
    run.font.bold = True

    paragraph = current_table.cell(1, 1).add_paragraph()  # год
    run = paragraph.add_run(str(filminfo[1]))
    run.font.name = 'Arial'
    run.font.size = Pt(10)

    paragraph = current_table.cell(1, 1).add_paragraph()  # страна
    run = paragraph.add_run(', '.join(filminfo[3]))
    run.font.name = 'Arial'
    run.font.size = Pt(10)

    paragraph = current_table.cell(1, 1).add_paragraph()  # режиссер
    run = paragraph.add_run('Режиссер: ' + filminfo[7])
    run.font.name = 'Arial'
    run.font.size = Pt(10)

    paragraph = current_table.cell(1, 1).add_paragraph()

    paragraph = current_table.cell(1, 1).add_paragraph()  # в главных ролях
    run = paragraph.add_run('В главных ролях: ')
    run.font.color.rgb = RGBColor(255, 102, 0)
    run.font.name = 'Arial'
    run.font.size = Pt(10)
    run = paragraph.add_run(', '.join(filminfo[8:18]))
    run.font.color.rgb = RGBColor(0, 0, 255)
    run.font.name = 'Arial'
    run.font.size = Pt(10)
    run.font.underline = True

    paragraph = current_table.cell(1, 1).add_paragraph()
    paragraph = current_table.cell(1, 1).add_paragraph()
    paragraph = current_table.cell(1, 1).add_paragraph()  # синопсис
    run = paragraph.add_run(filminfo[4])
    run.font.name = 'Arial'
    run.font.size = Pt(10)
    paragraph = current_table.cell(1, 1).add_paragraph()

    # запись постера в таблицу
    paragraph = current_table.cell(0, 0).paragraphs[1]
    run = paragraph.add_run()
    run.add_picture(image_to_file(filminfo[18]), width=Cm(7))


def write_all_films_to_docx(document, films: list, path: str):
    """Записывает информацию о фильмах в таблицы файла docx

    Args:
        document (_type_): Объект файла docx
        films (list): Список с информацией о фильмах
        path (str): Путь и имя для сохранения нового файла docx 

    """
    table_num = len(films)
    if table_num > 1:
        clone_first_table(document, table_num - 1)
    for i in tqdm(range(table_num), desc="Запись в таблицу...      "):
        current_table = document.tables[i]
        write_film_to_table(current_table, films[i])
    try:
        document.save(path)
        log.info(f'Файл "{path}" создан.')
    except PermissionError:
        log.warning(f"Ошибка! Нет доступа к файлу {path}. Список не сохранен.")
        raise Exception


def file_to_list(file: str):
    """Читает текстовый файл и возвращает список строк

    Args:
        file (str): Текстовый файл

    Raises:
        FileNotFoundError: файл не найден

    Returns:
        list: список строк из файла
    """
    if os.path.isfile(file):
        with open(file, 'r', encoding="utf-8") as f:
            list = [x.rstrip() for x in f]
        return list
    else:
        print(f'Файл {file} не найден.')
        raise FileNotFoundError


def write_tags_to_mp4(film: list, file_path: str):
    """Запись тегов в файл mp4.

    Args:
        film (list): Информация о фильме
        file_path (str): Путь к файлу mp4
    """
    video = MP4(file_path)
    video.delete()  # удаление всех тегов
    video["\xa9nam"] = film[0]  # title
    video["desc"] = film[4]  # description
    video["ldes"] = film[4]  # long description
    video["\xa9day"] = str(film[1])  # year
    video["covr"] = [MP4Cover(image_to_file(film[18]).getvalue(), imageformat=MP4Cover.FORMAT_PNG)]
    video.save()


def docx_to_pdf_libre(file_in):
    file_in_abs = os.path.abspath(file_in)
    dir_out_abs = os.path.split(file_in_abs)[0]
    soffice_path = "\"C:\Program Files\LibreOffice\program\soffice.exe\""
    command = f"{soffice_path} --headless --convert-to pdf --outdir {dir_out_abs} {file_in_abs}"
    os.system(command)


def main():
    from config import KINOPOISK_API_TOKEN
    parser = argparse.ArgumentParser(prog='Kinolist_Lib',
                                     description='Tool to create movie lists in docx format.',
                                     formatter_class=argparse.RawDescriptionHelpFormatter,
                                     epilog=textwrap.dedent("""\
                                     examples:
                                        Kinolist_Lib -m \"Terminator\" \"Terminator 2\" KP~319
                                        Kinolist_Lib -f movies.txt -o movies.docx
                                        Kinolist_Lib -t ./Terminator.mp4
                                        Kinolist_Lib -t c:\movies\Terminator.mp4
                                        Kinolist_Lib -t ./

                                        * Kinopoisk id can be set directly by placing tag KP~XXX in the title
                                        """))
    parser.add_argument("-ver", "--version", action="version", version=f"%(prog)s {LIB_VER}")
    parser.add_argument("-f", "--file", nargs=1, help="list of films in .txt format")
    parser.add_argument("-m", "--movie", nargs="+", help="list of films")
    parser.add_argument("-o", "--output", nargs=1, help="output file name (list.docx by default)")
    parser.add_argument("-s", "--shorten", action='store_true', help="shorten movie descriptions")
    parser.add_argument("-t", "--tag", nargs=1, help="write tags to mp4 file (or to all mp4 files in folder)")
    args = parser.parse_args()

    if args.output:
        output = args.output[0]
        output_dir, output_file_name = os.path.split(output)
        Path(output_dir).mkdir(parents=True, exist_ok=True)
        _, ext = os.path.splitext(output_file_name)
        if ext != ".docx":
            print("Output file must have .docx extension.")
            return
    else:
        output = "list.docx"

    if args.file:
        list = file_to_list((args.file[0]))
        print(f"Запрос из {args.file[0]} ({len(list)}): ", ", ".join(list))
        file_path = get_resource_path('template.docx')
        doc = Document(file_path)
        kp_codes = find_kp_id(list, KINOPOISK_API_TOKEN)
        if len(kp_codes[1]) != 0:
            for code in kp_codes[1]:
                log.warning(f"Фильм не найден, kinopoisk id: {code}")
        full_list = get_full_film_list(kp_codes[0], KINOPOISK_API_TOKEN, args.shorten)
        write_all_films_to_docx(doc, full_list, output)

    elif args.movie:
        film = args.movie
        kp_codes = find_kp_id(film, KINOPOISK_API_TOKEN)
        if len(kp_codes[1]) != 0:
            for code in kp_codes[1]:
                log.warning(f"Фильм не найден, kinopoisk id: {code}")
        if len(kp_codes[0]) == 0:
            log.warning("Фильмы не найдены.")
            return
        full_list = get_full_film_list(kp_codes[0], KINOPOISK_API_TOKEN, args.shorten)
        template_path = get_resource_path('template.docx')
        doc = Document(template_path)
        write_all_films_to_docx(doc, full_list, output)

    elif args.tag:
        path = args.tag[0]
        if os.path.isfile(path):
            _, mp4_file = os.path.split(path)
            name, ext = os.path.splitext(mp4_file)
            if ext != ".mp4":
                print("Can write tags only to mp4 files.")
                return
            name_list = []
            name_list.append(name)
            kp_id = find_kp_id(name_list, KINOPOISK_API_TOKEN)[0][0]
            film_info = get_film_info(kp_id, KINOPOISK_API_TOKEN)
            write_tags_to_mp4(film_info, path)
            log.info(f"Записан тег в файл: {mp4_file}")

        elif os.path.isdir(path):
            log.info(f"Поиск файлов mp4 в каталоге: {path}")
            mp4_files = glob.glob(path + '*.mp4')
            if len(mp4_files) < 1:
                log.warning(f'В каталоге "{path}" файлы mp4 не найдены.')
                return
            mp4_files_names = []
            for name in mp4_files:
                mp4_files_names.append(os.path.split(name)[1])
            for file in mp4_files_names:
                log.info(f"Найден файл: {file}")
            log.info(f"Всего найдено: {len(mp4_files)}")

            film_list = []
            for file in mp4_files:
                film_list.append(os.path.splitext(os.path.split(file)[1])[0])
            log.info("Поиск фильмов на kinopoisk.ru...")
            kp_id = find_kp_id(film_list, KINOPOISK_API_TOKEN)
            mp4_files_valid = []
            for index in range(len(mp4_files)):
                if film_list[index] not in kp_id[1]:
                    mp4_files_valid.append(mp4_files[index])
            full_films_list = get_full_film_list(kp_id[0], KINOPOISK_API_TOKEN)
            for index, film in enumerate(full_films_list):
                write_tags_to_mp4(film, mp4_files_valid[index])
                log.info(f"Записан тег в файл: {mp4_files_valid[index]}")
        else:
            log.error("Неверно указан путь.")


if __name__ == "__main__":
    main()
