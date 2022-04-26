import logging
import os
import re
import shutil
import sys
from copy import deepcopy

from aiogram import Bot, Dispatcher, executor, types
import requests
from docx import Document
from docx.shared import Cm, Pt, RGBColor
from kinopoisk_unofficial.kinopoisk_api_client import KinopoiskApiClient
from kinopoisk_unofficial.request.films.film_request import FilmRequest
from kinopoisk_unofficial.request.staff.staff_request import StaffRequest
from kinopoisk.movie import Movie
from PIL import Image
from docx2pdf import convert

import config

VER = '0.1.2'
TELEGRAM_API_TOKEN = config.TELEGRAM_API_TOKEN
KINOPOISK_API_TOKEN = config.KINOPOISK_API_TOKEN

# Configure logging
logging.basicConfig(level=logging.INFO)
log = logging.getLogger("Kinolist_Bot")

# Initialize bot and dispatcher
bot = Bot(token=TELEGRAM_API_TOKEN)
dp = Dispatcher(bot)


def is_api_ok(api):
    '''Проверка авторизации.'''
    try:
        api_client = KinopoiskApiClient(api)
        request = FilmRequest(328)
        api_client.films.send_film_request(request)
    except Exception:
        return False
    else:
        return True


def get_film_info(film_code, api):
    '''
    Получение информации о фильме с помощью kinopoisk_api_client.

            Элементы списка:
                0 - название фильма на русском языке
                1 - год
                2 - рейтинг Кинопоиска
                3 - список стран
                4 - описание
                5 - ссылка на постер
                6 - имя файла без расширения
                7 - режиссер
            8:17 - 10 актеров
    '''
    api_client = KinopoiskApiClient(api)
    request_staff = StaffRequest(film_code)
    response_staff = api_client.staff.send_staff_request(request_staff)
    stafflist = []
    for i in range(0, 11):  # загружаем 11 персоналий (режиссер + 10 актеров)
        if response_staff.items[i].name_ru == '':
            stafflist.append(response_staff.items[i].name_en)
        else:
            stafflist.append(response_staff.items[i].name_ru)

    request_film = FilmRequest(film_code)
    response_film = api_client.films.send_film_request(request_film)

    # с помощью регулярного выражения находим значение стран в кавычках ''
    countries = re.findall("'([^']*)'", str(response_film.film.countries))

    # имя файла
    filename = response_film.film.name_ru
    # очистка имени файла от запрещенных символов
    trtable = filename.maketrans('', '', '\/:*?"<>')
    filename = filename.translate(trtable)
    filmlist = [
        response_film.film.name_ru, response_film.film.year, response_film.film.rating_kinopoisk, countries,
        response_film.film.description, response_film.film.poster_url, filename
    ]
    return filmlist + stafflist


def write_film_to_table(current_table, filminfo, folder):
    '''Заполнение таблицы в файле docx.'''
    paragraph = current_table.cell(0, 1).paragraphs[0]  # название фильма + рейтинг
    run = paragraph.add_run(str(filminfo[0]) + ' - ' + 'Кинопоиск ' + str(filminfo[2]))
    run.font.name = 'Arial'
    run.font.size = Pt(11)
    run.font.bold = True

    paragraph = current_table.cell(1, 1).add_paragraph()  # год
    run = paragraph.add_run(str(filminfo[1]))
    run.font.name = 'Arial'
    run.font.size = Pt(10)

    paragraph = current_table.cell(1, 1).add_paragraph()  # страна
    run = paragraph.add_run(', '.join(filminfo[3]))
    run.font.name = 'Arial'
    run.font.size = Pt(10)

    paragraph = current_table.cell(1, 1).add_paragraph()  # режиссер
    run = paragraph.add_run('Режиссер: ' + filminfo[7])
    run.font.name = 'Arial'
    run.font.size = Pt(10)

    paragraph = current_table.cell(1, 1).add_paragraph()

    paragraph = current_table.cell(1, 1).add_paragraph()  # в главных ролях
    run = paragraph.add_run('В главных ролях: ')
    run.font.color.rgb = RGBColor(255, 102, 0)
    run.font.name = 'Arial'
    run.font.size = Pt(10)
    run = paragraph.add_run(', '.join(filminfo[8:]))
    run.font.color.rgb = RGBColor(0, 0, 255)
    run.font.name = 'Arial'
    run.font.size = Pt(10)
    run.font.underline = True

    paragraph = current_table.cell(1, 1).add_paragraph()
    paragraph = current_table.cell(1, 1).add_paragraph()
    paragraph = current_table.cell(1, 1).add_paragraph()  # синопсис
    run = paragraph.add_run(filminfo[4])
    run.font.name = 'Arial'
    run.font.size = Pt(10)
    paragraph = current_table.cell(1, 1).add_paragraph()

    # загрузка постера
    image_url = filminfo[5]
    file_path = './' + folder + '/covers/' + str(filminfo[6] + '.jpg')
    resp = requests.get(image_url, stream=True)
    if resp.status_code == 200:
        resp.raw.decode_content = True
        with open(file_path, 'wb') as f:  # открываем файл для бинарной записи
            shutil.copyfileobj(resp.raw, f)
    else:
        log.warning(f'Не удалось загрузить постер ({image_url})')

    # изменение размера постера
    image = Image.open(file_path)
    width, height = image.size
    # обрезка до соотношения сторон 1x1.5
    if width > (height / 1.5):
        image = image.crop((((width - height / 1.5) / 2), 0, ((width - height / 1.5) / 2) + height / 1.5, height))
    image.thumbnail((360, 540))
    rgb_image = image.convert('RGB')  # для исправление возможной ошибки "OSError: cannot write mode RGBA as JPEG"
    rgb_image.save(file_path)

    # запись постера в таблицу
    paragraph = current_table.cell(0, 0).paragraphs[1]
    run = paragraph.add_run()
    run.add_picture(file_path, width=Cm(7))


def copy_table_after(table, paragraph):
    '''Копирование таблицы в указанный параграф.'''
    tbl, p = table._tbl, paragraph._p
    new_tbl = deepcopy(tbl)
    p.addnext(new_tbl)


def clone_first_table(document: Document, num):
    '''Клонирует первую таблицу в документе num раз.'''
    template = document.tables[0]
    paragraph = document.paragraphs[0]
    for i in range(num):
        copy_table_after(template, paragraph)
        paragraph = document.add_paragraph()


def get_resource_path(relative_path):
    '''
    Определение пути для запуска из автономного exe файла.

    Pyinstaller cоздает временную папку, путь в _MEIPASS.
    '''
    try:
        base_path = sys._MEIPASS
    except Exception:
        base_path = os.path.abspath(".")
    return os.path.join(base_path, relative_path)


@dp.message_handler(commands=['start', 'help'])
async def send_welcome(message: types.Message):
    """
    This handler will be called when user sends `/start` or `/help` command
    """
    log.info(f"Start (chat_id: {message.chat.id})")
    await message.reply("Привет, я Кinolist Bot!\nОтправьте мне список фильмов, и я пришлю его в формате pdf.")


@dp.message_handler(commands=['lisa', 'Lisa'])
async def send_heart(message: types.Message):
    await message.reply_sticker("CAACAgIAAxkBAAEEjSZiZXLQqPDFY70qC0m9PPH2AAEJjfgAAjIAA-Sgzgd7_cFVbY2YfiQE")
    log.info("Отправлен стикер")


@dp.message_handler()
async def reply(message: types.Message):
    if not is_api_ok(KINOPOISK_API_TOKEN):
        log.warning("API error.")
        await message.reply("Ой, что-то сломалось!((\n" + "(API error)")
        return
    chat_id = str(message.chat.id)
    log.info(f"Start list generate (chat_id: {chat_id})")
    if os.path.isdir(chat_id):
        log.info(f"Папка {chat_id} обнаружена")
        await message.reply("Подождите, я все еще работаю!")
        return
    film_list = message.text.split('\n')
    film_list = list(filter(None, film_list))
    log.info(film_list)
    film_codes = []
    film_not_found = []
    for film in film_list:
        try:
            found_films = Movie.objects.search(film)
        except Exception:
            log.info(f'{film} не найден')
            film_not_found.append(film)
            continue
        else:
            if len(found_films) < 1:
                log.info(f'{film} не найден')
                film_not_found.append(film)
                continue
            id = str(found_films[0].id)
            log.info(f'Найден фильм: {found_films[0]}, kinopoisk id: {id}')
            film_codes.append(id)
    if len(film_not_found) > 0:
        log.info(f'Не найдено: {", ".join(film_not_found)}')
    if len(film_codes) < 1:
        await message.reply("Ой, ничего не найдено!")
        return
    full_films_list = []
    for film_code in film_codes:
        try:
            film_info = get_film_info(film_code, KINOPOISK_API_TOKEN)
            full_films_list.append(film_info)
        except Exception:
            log.warning(f'{film_code} - ошибка')
        else:
            continue
    if len(full_films_list) < 1:
        await message.reply("Ни один фильм не найден!")
        return
    file_path = get_resource_path('template.docx')
    try:
        doc = Document(file_path)
    except Exception:
        log.warning('Не найден шаблон "template.docx". Список не создан.')
        await message.reply("Ой, что-то сломалось!((")
        return
    table_num = len(full_films_list)
    if table_num > 1:
        clone_first_table(doc, table_num - 1)
    if not os.path.isdir(chat_id):
        os.mkdir('./' + chat_id)
        os.mkdir('./' + chat_id + '/covers/')
    for i in range(table_num):
        current_table = doc.tables[i]
        write_film_to_table(current_table, full_films_list[i], chat_id)
        log.info(f'{full_films_list[i][0]} - ок')
    try:
        doc.save('./' + chat_id + '/list.docx')
    except PermissionError:
        log.warning("Ошибка! Нет доступа к файлу list.docx. Список не создан.")
        await message.reply("Ой, что-то сломалось!((")
    convert('./' + chat_id + "/list.docx", './' + chat_id + "/list.pdf")
    with open('./' + chat_id + "/list.pdf", 'rb') as pdf:
        if len(film_not_found) > 0:
            text = "Список готов!" + "\n" + "Правда, вот эти фильмы не смог найти:" + "\n" + "\n".join(film_not_found)
            await message.reply_document(pdf, caption=text)
        else:
            await message.reply_document(pdf, caption='Список готов!')
    shutil.rmtree('./' + chat_id)
    return


if __name__ == '__main__':
    executor.start_polling(dp, skip_updates=True)
